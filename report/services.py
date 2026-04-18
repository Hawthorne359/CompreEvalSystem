"""
报表导出服务：统一查询、模板渲染、文件响应。
"""
from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
from collections import defaultdict
from pathlib import Path
from urllib.parse import quote

from django.db.models import Q
from django.http import HttpResponse

from eval.models import EvalIndicator
from scoring.models import ArbitrationRecord, ImportedScoreDetail, ReviewAssignment, ScoreRecord
from scoring.services import get_indicator_final_score
from submission.models import Evidence, StudentSubmission, SubmissionAnswer
from users.permissions import user_level_at_least




def _build_soft_undefined():
    """构造宽松 Undefined 类型，避免模板缺字段时中断导出。"""
    from jinja2 import ChainableUndefined

    class SoftUndefined(ChainableUndefined):
        def __str__(self):
            return ''

        def __bool__(self):
            return False

        def __int__(self):
            return 0

        def __float__(self):
            return 0.0

        def __add__(self, other):
            return 0

        __radd__ = __add__
        __sub__ = __add__
        __rsub__ = __add__
        __mul__ = __add__
        __rmul__ = __add__
        __truediv__ = __add__
        __rtruediv__ = __add__
        __floordiv__ = __add__
        __rfloordiv__ = __add__
        __mod__ = __add__
        __rmod__ = __add__
        __pow__ = __add__
        __rpow__ = __add__

        def __lt__(self, other):
            return False

        __le__ = __lt__
        __gt__ = __lt__
        __ge__ = __lt__

    return SoftUndefined


FIELD_GROUPS = [
    {'id': 'user', 'label': '用户字段', 'order': 10},
    {'id': 'season_project', 'label': '周期与项目字段', 'order': 20},
    {'id': 'submission', 'label': '提交与成绩字段', 'order': 30},
    {'id': 'indicators', 'label': '指标拆分字段', 'order': 40},
    {'id': 'evidence', 'label': '佐证字段', 'order': 50},
    {'id': 'sensitive', 'label': '敏感字段', 'order': 60},
]

FIELD_VIEW_MODE_TEMPLATE_COMMON = 'template_common_first'
FIELD_VIEW_MODE_ADVANCED_ALL = 'advanced_all_fields'
_EXCEL_COLUMN_RE = re.compile(r'^[A-Za-z]{1,3}$')
_CELL_RE = re.compile(r'^[A-Za-z]{1,3}[1-9]\d*$')
SUPPORTED_GROUP_BY = {'', 'none', 'class', 'major', 'department'}
SUPPORTED_GROUP_FILE_MODE = {'single_file', 'per_group_file', 'per_student_file'}

COMMON_BASE_KEYS = {
    'season_academic_year',
    'department',
    'major',
    'class_name',
    'real_name',
    'student_no',
    'season_name',
    'season_semester',
}
DEFAULT_LONG_FORM_BASE_KEYS = [
    'season_academic_year',
    'department',
    'major',
    'class_name',
    'real_name',
    'student_no',
]

BASE_FIELDS = [
    {'key': 'rank', 'label': '排名', 'category_id': 'submission', 'order': 1, 'is_common': False},
    {'key': 'student_no', 'label': '学号', 'category_id': 'user', 'order': 2},
    {'key': 'username', 'label': '用户名', 'category_id': 'user', 'order': 3},
    {'key': 'real_name', 'label': '姓名', 'category_id': 'user', 'order': 4},
    {'key': 'gender', 'label': '性别', 'category_id': 'user', 'order': 5},
    {'key': 'department', 'label': '院系', 'category_id': 'user', 'order': 6},
    {'key': 'major', 'label': '专业', 'category_id': 'user', 'order': 7},
    {'key': 'class_name', 'label': '班级', 'category_id': 'user', 'order': 8},
    {'key': 'class_grade', 'label': '年级', 'category_id': 'user', 'order': 9},
    {'key': 'project_name', 'label': '测评项目', 'category_id': 'season_project', 'order': 10, 'is_common': True},
    {'key': 'project_status', 'label': '项目状态', 'category_id': 'season_project', 'order': 11},
    {'key': 'project_start_time', 'label': '项目开始时间', 'category_id': 'season_project', 'order': 12},
    {'key': 'project_end_time', 'label': '项目结束时间', 'category_id': 'season_project', 'order': 13},
    {'key': 'project_review_end_time', 'label': '项目评定截止时间', 'category_id': 'season_project', 'order': 14},
    {'key': 'season_name', 'label': '测评周期', 'category_id': 'season_project', 'order': 15},
    {'key': 'season_academic_year', 'label': '学年', 'category_id': 'season_project', 'order': 16},
    {'key': 'season_semester', 'label': '学期', 'category_id': 'season_project', 'order': 17},
    {'key': 'submission_id', 'label': '提交ID', 'category_id': 'submission', 'order': 18, 'is_common': False},
    {'key': 'submission_status', 'label': '提交状态', 'category_id': 'submission', 'order': 19},
    {'key': 'submitted_at', 'label': '提交时间', 'category_id': 'submission', 'order': 20},
    {'key': 'remark', 'label': '提交备注', 'category_id': 'submission', 'order': 21},
    {'key': 'final_score', 'label': '总分', 'category_id': 'submission', 'order': 22, 'is_common': True},
    {'key': 'scope_valid_submission_count', 'label': '参加测评人数', 'category_id': 'submission', 'order': 22.1, 'is_common': True},
    {'key': 'scope_missing_submission_count', 'label': '未参加测评人数', 'category_id': 'submission', 'order': 22.2, 'is_common': True},
    {'key': 'scope_total_submission_count', 'label': '应参加总人数', 'category_id': 'submission', 'order': 22.3},
    {'key': 'reviewer_signatures_all', 'label': '测评小组成员签名（全部）', 'category_id': 'submission', 'order': 22.4, 'is_common': True},
    {'key': 'reviewer_signatures_teachers', 'label': '测评小组成员签名（评审老师）', 'category_id': 'submission', 'order': 22.5},
    {'key': 'reviewer_signatures_assistants', 'label': '测评小组成员签名（学生助理）', 'category_id': 'submission', 'order': 22.6},
    {'key': 'global_evidence_count', 'label': '全局佐证数量', 'category_id': 'evidence', 'order': 23},
    {'key': 'global_evidence_names', 'label': '全局佐证名称列表', 'category_id': 'evidence', 'order': 24},
    {'key': 'global_evidence_urls', 'label': '全局佐证链接列表', 'category_id': 'evidence', 'order': 25},
    {'key': 'phone', 'label': '手机号', 'category_id': 'sensitive', 'order': 26, 'is_sensitive': True},
    {'key': 'employee_no', 'label': '工号', 'category_id': 'sensitive', 'order': 27, 'is_sensitive': True},
]


def _safe_file_name(file_name: str) -> str:
    sanitized = (file_name or 'report').replace('"', '').replace('\n', '').replace('\r', '')
    return sanitized


def build_download_response(content: bytes, content_type: str, file_name: str) -> HttpResponse:
    safe_name = _safe_file_name(file_name)
    encoded = quote(safe_name)
    resp = HttpResponse(content, content_type=content_type)
    resp['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded}"
    return resp


def _resolve_submission_queryset(project, user, filters: dict | None = None, *, include_unscored: bool = False):
    filters = filters or {}
    submissions = StudentSubmission.objects.filter(project=project).select_related(
        'user',
        'user__class_obj',
        'user__class_obj__major',
        'user__department',
    )
    if not include_unscored:
        submissions = submissions.exclude(final_score__isnull=True)
    current_level = user.current_role.level if user.current_role else -1
    if current_level == 3 and user.department_id:
        submissions = submissions.filter(user__department_id=user.department_id)
    elif not user_level_at_least(user, 3):
        from users.models import UserRole
        scope_class_ids = list(
            UserRole.objects.filter(user=user, scope_type='class').values_list('scope_id', flat=True)
        )
        submissions = submissions.filter(user__class_obj_id__in=scope_class_ids) if scope_class_ids else submissions.none()
    if user_level_at_least(user, 5):
        if filters.get('department_id'):
            submissions = submissions.filter(user__department_id=filters['department_id'])
    if filters.get('major_id'):
        submissions = submissions.filter(user__class_obj__major_id=filters['major_id'])
    if filters.get('class_id'):
        submissions = submissions.filter(user__class_obj_id=filters['class_id'])
    if filters.get('search'):
        kw = filters['search'].strip()
        if kw:
            submissions = submissions.filter(
                Q(user__name__icontains=kw)
                | Q(user__username__icontains=kw)
                | Q(user__student_no__icontains=kw)
            )
    return submissions


def resolve_report_queryset(project, user, filters: dict | None = None):
    return _resolve_submission_queryset(project, user, filters=filters, include_unscored=False).order_by('-final_score', 'id')


def _collect_leaf_indicators(project):
    indicators = list(
        EvalIndicator.objects.filter(project=project, score_source__in=('self', 'import', 'reviewer')).select_related('parent').order_by('order', 'id')
    )
    used_as_parent = set(
        EvalIndicator.objects.filter(project=project, parent_id__isnull=False).values_list('parent_id', flat=True)
    )
    return [ind for ind in indicators if ind.id not in used_as_parent]


def _collect_parent_indicators(project):
    """Return all indicators that have children (i.e. appear as parent_id of another indicator).
    score_detail stores aggregate scores for ALL parent nodes regardless of score_source,
    so we must not restrict to score_source='children' only.
    """
    used_as_parent = set(
        EvalIndicator.objects.filter(project=project, parent_id__isnull=False)
        .values_list('parent_id', flat=True)
    )
    return list(
        EvalIndicator.objects.filter(project=project, id__in=used_as_parent)
        .select_related('parent')
        .order_by('order', 'id')
    )


def _category_index():
    return {g['id']: g for g in FIELD_GROUPS}


def _module_order(module_key: str) -> int:
    """
    Derive a stable sort order from a module key like 'F1', 'F2', 'A', 'B', etc.
    Numeric suffix determines fine-grained ordering within the same letter group.
    """
    key = (module_key or '').strip().upper()
    # Extract leading letter(s) and optional number, e.g. 'F1' -> letter='F', num=1
    m = re.match(r'^([A-Z]+)(\d*)$', key)
    if not m:
        return 9990
    letter_part = m.group(1)
    num_part = int(m.group(2)) if m.group(2) else 0
    # Compute letter base dynamically from ASCII values so any letter prefix works.
    # Treats multi-letter prefixes as base-26 numbers: A=1, B=2, ..., Z=26.
    # e.g. 'A'->1, 'F'->6, 'Z'->26, 'AA'->27, 'AB'->28
    letter_base = 0
    for c in letter_part:
        letter_base = letter_base * 26 + (ord(c) - ord('A') + 1)
    return letter_base * 10000 + num_part


def _field_type_order(split_type: str) -> int:
    # 常用字段优先：自评分 -> 过程记录 -> 最终分
    order_map = {
        'self_score': 10,
        'process_record': 20,
        'final_adopted_score': 30,
        'final_adopted': 35,
        'imported_score': 100,
        'reviewer_score': 110,
        'arbitration_score': 120,
        'evidence_count': 130,
        'evidence_names': 140,
        'evidence_urls': 150,
    }
    return order_map.get(split_type or '', 999)


def _extract_indicator_index(indicator_key: str) -> int:
    """
    从指标编码中提取序号（如 A1 -> 1），用于稳定排序。
    无法提取时返回较大值，避免打乱已识别项顺序。
    """
    if not indicator_key:
        return 9999
    m = re.match(r'^[A-Za-z](\d+)$', str(indicator_key).strip())
    if not m:
        return 9999
    try:
        return int(m.group(1))
    except (TypeError, ValueError):
        return 9999


def _extract_module_key(indicator):
    candidates = [
        (getattr(indicator, 'category', '') or '').strip(),
        (getattr(indicator, 'name', '') or '').strip(),
        ((indicator.parent.name if getattr(indicator, 'parent_id', None) and getattr(indicator, 'parent', None) else '') or '').strip(),
    ]
    for text in candidates:
        m = re.match(r'^([A-Za-z])\d*', text)
        if m:
            return m.group(1).upper()
    return 'Z'


def _get_root_parent(indicator, id_map=None):
    """
    Walk up to the true top-level (root) indicator (the one with no parent).
    Supports arbitrary depth; uses id_map for efficient traversal when provided.
    """
    cur = indicator
    for _ in range(20):  # guard against cycles
        if not getattr(cur, 'parent_id', None):
            return cur
        parent = None
        if id_map is not None:
            parent = id_map.get(cur.parent_id)
        if parent is None:
            parent = getattr(cur, 'parent', None)
        if parent is None:
            return cur
        cur = parent
    return cur


def _extract_module_code(indicator):
    """
    Extract the full module code (e.g. 'F1', 'F2') from an indicator's category field.
    Falls back to an empty string if not found.
    """
    category = (getattr(indicator, 'category', '') or '').strip()
    m = re.match(r'^([A-Za-z]\d+)', category)
    if m:
        return m.group(1).upper()
    return ''


def _extract_indicator_code(indicator):
    candidates = [
        (getattr(indicator, 'name', '') or '').strip(),
        (getattr(indicator, 'category', '') or '').strip(),
    ]
    for text in candidates:
        m = re.match(r'^([A-Za-z]\d+)', text)
        if m:
            return m.group(1).upper()
    return f'IND{indicator.id}'


def _build_field_tree(fields):
    module_map = {}
    for field in fields:
        if field.get('category_id') != 'indicators':
            continue
        module_key = field.get('module_key') or 'Z'
        module_node = module_map.setdefault(module_key, {
            'module_key': module_key,
            'module_label': field.get('module_label') or module_key,
            'module_code': field.get('module_code') or '',
            'module_order': field.get('module_order', 999),
            'children_map': {},
        })
        # Prefer non-empty module_code once we see it
        if not module_node['module_code'] and field.get('module_code'):
            module_node['module_code'] = field['module_code']
        ind_key = str(field.get('indicator_id') or field.get('group_label') or field.get('key'))
        ind_node = module_node['children_map'].setdefault(ind_key, {
            'indicator_id': field.get('indicator_id'),
            'indicator_key': field.get('indicator_key') or ind_key,
            'indicator_label': field.get('group_label') or field.get('indicator_name') or ind_key,
            'indicator_path': field.get('indicator_path') or [],
            'indicator_order': field.get('indicator_order', 999999),
            'fields': [],
        })
        ind_node['fields'].append({
            'key': field.get('key'),
            'label': field.get('label'),
            'split_type': field.get('split_type'),
            'is_common': bool(field.get('is_common')),
            'field_type_order': field.get('field_type_order', 999),
        })
    tree = []
    for module in module_map.values():
        indicators = list(module['children_map'].values())
        for ind in indicators:
            ind['fields'].sort(key=lambda x: (x.get('field_type_order', 999), x.get('label', '')))
        indicators.sort(key=lambda x: (x.get('indicator_order', 999999), x.get('indicator_label', '')))
        tree.append({
            'module_key': module['module_key'],
            'module_label': module['module_label'],
            'module_code': module['module_code'],
            'module_order': module['module_order'],
            'children': indicators,
        })
    tree.sort(key=lambda x: (x.get('module_order', 999), x.get('module_label', '')))
    return tree


def _build_indicator_field_tree(fields, id_map):
    """
    Build a truly nested indicator tree that mirrors the DB indicator hierarchy.
    Each node shape:
      {
        'id': int,
        'name': str,
        'category': str,        # e.g. 'F1', 'A1', '' for intermediate nodes
        'score_source': str,
        'fields': [...],        # field dicts for THIS indicator (sorted by field_type_order)
        'children': [...],      # recursive child nodes
      }
    Fields are attached to the node whose indicator_id matches.
    Nodes without any fields AND no children that have fields are still shown so the
    full hierarchy is visible; they just have an empty fields list.
    """
    # Build field_map: indicator_id -> sorted list of field dicts
    field_map: dict = {}
    for f in fields:
        if f.get('category_id') != 'indicators':
            continue
        iid = f.get('indicator_id')
        if iid is not None:
            field_map.setdefault(iid, []).append(f)
    for iid, flist in field_map.items():
        flist.sort(key=lambda x: (x.get('field_type_order', 999), x.get('label', '')))

    def _total_fields(ind_id):
        """Recursively count all fields reachable from this node."""
        total = len(field_map.get(ind_id, []))
        for child in [x for x in id_map.values() if x.parent_id == ind_id]:
            total += _total_fields(child.id)
        return total

    def build_node(ind):
        child_inds = sorted(
            [x for x in id_map.values() if x.parent_id == ind.id],
            key=lambda x: (x.order, x.id),
        )
        raw_fields = field_map.get(ind.id, [])
        return {
            'id': ind.id,
            'name': ind.name,
            'category': ind.category or '',
            'score_source': ind.score_source,
            'total_field_count': _total_fields(ind.id),
            'fields': [
                {
                    'key': f['key'],
                    'label': f['label'],
                    'split_type': f['split_type'],
                    'is_common': bool(f.get('is_common')),
                    'field_type_order': f.get('field_type_order', 999),
                }
                for f in raw_fields
            ],
            'children': [build_node(c) for c in child_inds],
        }

    root_indicators = sorted(
        [ind for ind in id_map.values() if not ind.parent_id],
        key=lambda x: (x.order, x.id),
    )
    return [build_node(r) for r in root_indicators]


def _build_default_word_presets(fields):
    """
    生成可直接套用的占位符预设（@token 模式）。
    """
    field_map = {f.get('key'): f for f in fields}
    placeholders = []
    exists = set()

    def _append(field_key):
        if not field_key or field_key in exists or field_key not in field_map:
            return
        placeholders.append({
            'placeholder': field_key,
            'field_key': field_key,
            'label': field_map[field_key].get('label') or field_key,
        })
        exists.add(field_key)

    for key in DEFAULT_LONG_FORM_BASE_KEYS:
        _append(key)

    indicator_fields = {}
    for f in fields:
        if f.get('category_id') != 'indicators':
            continue
        iid = f.get('indicator_id')
        if iid is None:
            continue
        indicator_fields.setdefault(iid, []).append(f)
    ordered_indicator_items = sorted(
        indicator_fields.items(),
        key=lambda item: min([(x.get('order', 999999), x.get('key', '')) for x in item[1]] or [(999999, '')]),
    )
    for _, group in ordered_indicator_items:
        group_map = {x.get('split_type'): x for x in group}
        record_only = any(bool(x.get('is_record_only')) for x in group)
        if record_only:
            _append((group_map.get('self_score') or {}).get('key'))
            process_field = group_map.get('process_record') or {}
            if bool(process_field.get('require_process_record')):
                _append(process_field.get('key'))
            continue
        _append((group_map.get('self_score') or {}).get('key'))
        _append((group_map.get('process_record') or {}).get('key'))
        _append((group_map.get('final_adopted_score') or {}).get('key'))

    return [
        {
            'id': 'long_form_default',
            'label': '长表默认模板（基础信息 + 指标常用字段）',
            'token_mode': 'prefix_token',
            'token_prefix': '@',
            'word_placeholders': placeholders,
        }
    ]


def _build_indicator_path_labels(indicator, id_map):
    """
    Build the full ancestor name path for an indicator (excluding the root).
    Returns (prefix_label, indicator_path_list).
    e.g. for 必修课门次 whose ancestors are 智育评分->B1:
        prefix_label = 'B1/必修课门次'
        indicator_path = ['智育评分', 'B1', '必修课门次']
    """
    path = []
    cur = indicator
    for _ in range(20):
        path.insert(0, cur.name)
        if not getattr(cur, 'parent_id', None):
            break
        parent = id_map.get(cur.parent_id) if id_map else getattr(cur, 'parent', None)
        if parent is None:
            break
        cur = parent
    # path[0] is the root name, rest are intermediate + leaf
    # prefix_label = everything except the root, joined with '/'
    if len(path) > 1:
        prefix_label = '/'.join(path[1:])
    else:
        prefix_label = path[0] if path else indicator.name
    return prefix_label, path


def get_project_export_catalog(project, view_mode=FIELD_VIEW_MODE_TEMPLATE_COMMON, user_common_keys=None):
    user_common_set = set(str(k).strip() for k in (user_common_keys or []) if str(k).strip())
    # Pre-load ALL indicators for this project so we can walk up any depth
    all_indicators = list(
        EvalIndicator.objects.filter(project=project).order_by('order', 'id')
    )
    id_map = {ind.id: ind for ind in all_indicators}
    # Restore parent references using id_map (since we didn't select_related deeply)
    for ind in all_indicators:
        if ind.parent_id and ind.parent_id in id_map and not getattr(ind, '_parent_loaded', False):
            ind.parent = id_map[ind.parent_id]
            ind._parent_loaded = True

    leaf_indicators = _collect_leaf_indicators(project)
    # Re-attach parents from id_map for leaf indicators
    for ind in leaf_indicators:
        if ind.parent_id and ind.parent_id in id_map:
            ind.parent = id_map[ind.parent_id]

    fields = []
    category_meta = _category_index()
    for field in BASE_FIELDS:
        f = dict(field)
        f['source'] = 'system'
        f['category_label'] = category_meta.get(f.get('category_id'), {}).get('label', '')
        f['group_order'] = category_meta.get(f.get('category_id'), {}).get('order', 999)
        system_common = bool(f.get('is_common', f.get('key') in COMMON_BASE_KEYS))
        user_common = f.get('key') in user_common_set
        f['is_system_common'] = system_common
        f['is_user_common'] = user_common
        f['is_common'] = bool(system_common or user_common)
        f['priority'] = 10 if f['is_common'] else 100
        fields.append(f)
    for ind in leaf_indicators:
        prefix_label, indicator_path = _build_indicator_path_labels(ind, id_map)
        root = _get_root_parent(ind, id_map)
        module_code = _extract_module_code(root)
        # Use module_code (e.g. 'F1') as module_key for correct grouping;
        # fall back to single-letter extraction only when no numeric code is found.
        module_key = module_code if module_code else _extract_module_key(root)
        indicator_key = _extract_indicator_code(ind)
        module_label = root.name
        module_order = _module_order(module_key)
        indicator_index = _extract_indicator_index(indicator_key)
        base_order = module_order * 100000 + indicator_index * 100 + ind.id
        review_enabled_for_record_only = bool(getattr(ind, 'record_only_requires_review', False))
        has_review_stage = not ind.is_record_only or review_enabled_for_record_only
        split_fields = []
        # Build field list based on score_source so only meaningful fields are shown.
        # All types share: process_record, evidence_* (gated by is_record_only below).
        if ind.score_source == 'self':
            split_fields.append((f'ind_{ind.id}_self_score', f'{prefix_label}-自评分', 'self_score'))
            if has_review_stage:
                split_fields.extend([
                    (f'ind_{ind.id}_reviewer_score', f'{prefix_label}-评审分', 'reviewer_score'),
                    (f'ind_{ind.id}_arbitration_score', f'{prefix_label}-仲裁分', 'arbitration_score'),
                    (f'ind_{ind.id}_final_adopted_score', f'{prefix_label}-确认分', 'final_adopted_score'),
                ])
        elif ind.score_source == 'import':
            split_fields.extend([
                (f'ind_{ind.id}_imported_score', f'{prefix_label}-导入分', 'imported_score'),
                (f'ind_{ind.id}_final_adopted_score', f'{prefix_label}-确认分', 'final_adopted_score'),
            ])
        elif ind.score_source == 'reviewer':
            if has_review_stage:
                split_fields.extend([
                    (f'ind_{ind.id}_reviewer_score', f'{prefix_label}-评审分', 'reviewer_score'),
                    (f'ind_{ind.id}_arbitration_score', f'{prefix_label}-仲裁分', 'arbitration_score'),
                    (f'ind_{ind.id}_final_adopted_score', f'{prefix_label}-确认分', 'final_adopted_score'),
                ])
        else:
            # Fallback for any unknown score_source: expose all score fields
            split_fields.append((f'ind_{ind.id}_self_score', f'{prefix_label}-自评分', 'self_score'))
            if has_review_stage:
                split_fields.extend([
                    (f'ind_{ind.id}_imported_score', f'{prefix_label}-导入分', 'imported_score'),
                    (f'ind_{ind.id}_reviewer_score', f'{prefix_label}-评审分', 'reviewer_score'),
                    (f'ind_{ind.id}_arbitration_score', f'{prefix_label}-仲裁分', 'arbitration_score'),
                    (f'ind_{ind.id}_final_adopted_score', f'{prefix_label}-确认分', 'final_adopted_score'),
                ])
        split_fields.extend([
            (f'ind_{ind.id}_process_record', f'{prefix_label}-过程记录', 'process_record'),
            (f'ind_{ind.id}_evidence_count', f'{prefix_label}-佐证数量', 'evidence_count'),
            (f'ind_{ind.id}_evidence_names', f'{prefix_label}-佐证名称列表', 'evidence_names'),
            (f'ind_{ind.id}_evidence_urls', f'{prefix_label}-佐证链接列表', 'evidence_urls'),
        ])
        for idx, (key, label, split_type) in enumerate(split_fields):
            system_common = False
            if split_type == 'final_adopted_score':
                system_common = True
            elif split_type == 'self_score' and ind.score_source == 'self':
                system_common = True
            elif split_type == 'imported_score' and ind.score_source == 'import':
                system_common = True
            elif split_type == 'process_record' and ind.score_source == 'self' and ind.require_process_record:
                system_common = True
            user_common = key in user_common_set
            is_common = bool(system_common or user_common)
            fields.append({
                'key': key,
                'label': label,
                'source': ind.score_source,
                'indicator_id': ind.id,
                'indicator_name': ind.name,
                'indicator_key': indicator_key,
                'split_type': split_type,
                'field_type_order': _field_type_order(split_type),
                'category_id': 'indicators',
                'category_label': category_meta['indicators']['label'],
                'group_label': prefix_label,
                'module_key': module_key,
                'module_code': module_code,
                'module_label': module_label,
                'module_order': module_order,
                'indicator_path': indicator_path,
                'indicator_order': base_order,
                'is_record_only': bool(ind.is_record_only),
                'record_only_requires_review': review_enabled_for_record_only,
                'require_process_record': bool(ind.require_process_record),
                'is_system_common': system_common,
                'is_user_common': user_common,
                'is_common': is_common,
                'priority': 20 if is_common else 200,
                'group_order': category_meta['indicators']['order'],
                'order': base_order + idx,
            })
    parent_indicators = _collect_parent_indicators(project)
    for ind in parent_indicators:
        # Re-attach parent from id_map for multi-level support
        if ind.parent_id and ind.parent_id in id_map:
            ind.parent = id_map[ind.parent_id]
        root = _get_root_parent(ind, id_map)
        module_code = _extract_module_code(root)
        module_key = module_code if module_code else _extract_module_key(root)
        indicator_key = _extract_indicator_code(ind)
        name_label = ind.name
        module_label = root.name
        module_order = _module_order(module_key)
        indicator_index = _extract_indicator_index(indicator_key)
        base_order = module_order * 100000 + indicator_index * 100 + ind.id
        agg_key = f'agg_{ind.id}'
        fields.append({
            'key': agg_key,
            'label': f'{name_label}-汇总分',
            'source': 'children',
            'indicator_id': ind.id,
            'indicator_name': ind.name,
            'indicator_key': indicator_key,
            'split_type': 'agg_score',
            'field_type_order': 5,
            'category_id': 'indicators',
            'category_label': category_meta['indicators']['label'],
            'group_label': name_label,
            'module_key': module_key,
            'module_code': module_code,
            'module_label': module_label,
            'module_order': module_order,
            'indicator_path': _build_indicator_path_labels(ind, id_map)[1],
            'indicator_order': base_order,
            'is_record_only': False,
            'record_only_requires_review': False,
            'require_process_record': False,
            'is_system_common': True,
            'is_user_common': agg_key in user_common_set,
            'is_common': True,
            'priority': 15,
            'group_order': category_meta['indicators']['order'],
            'order': base_order,
        })
    fields.sort(key=lambda x: (x.get('priority', 999), x.get('group_order', 999), x.get('order', 999999), x.get('label', ''), x.get('key', '')))
    if view_mode == FIELD_VIEW_MODE_TEMPLATE_COMMON:
        selected_fields = [f for f in fields if f.get('is_common')]
    else:
        selected_fields = list(fields)
    return {
        'field_version': 2,
        'common_profile_version': 1,
        'field_view_mode': view_mode,
        'field_groups': FIELD_GROUPS,
        'fields': selected_fields,
        'all_fields': fields,
        'field_tree': _build_indicator_field_tree(fields, id_map),
        'presets': _build_default_word_presets(fields),
        'user_common_keys': sorted(user_common_set),
    }


def get_project_export_fields(project):
    return get_project_export_catalog(project)['fields']


def build_export_rows(project, submissions):
    leaf_indicators = _collect_leaf_indicators(project)
    parent_indicators = _collect_parent_indicators(project)
    indicator_ids = [ind.id for ind in leaf_indicators]
    indicators = {ind.id: ind for ind in EvalIndicator.objects.filter(id__in=indicator_ids)}
    submission_ids = [sub.id for sub in submissions]

    answer_map = {
        (row['submission_id'], row['indicator_id']): row
        for row in SubmissionAnswer.objects.filter(submission_id__in=submission_ids, indicator_id__in=indicator_ids)
        .values('submission_id', 'indicator_id', 'self_score', 'process_record')
    }
    imported_map = {
        (row['submission_id'], row['indicator_id']): row['score']
        for row in ImportedScoreDetail.objects.filter(submission_id__in=submission_ids, indicator_id__in=indicator_ids)
        .values('submission_id', 'indicator_id', 'score')
    }
    arbitration_map = {
        (row['submission_id'], row['indicator_id']): row['score']
        for row in ArbitrationRecord.objects.filter(submission_id__in=submission_ids, indicator_id__in=indicator_ids)
        .values('submission_id', 'indicator_id', 'score')
    }
    review_map = {}
    for row in (
        ScoreRecord.objects.filter(submission_id__in=submission_ids, indicator_id__in=indicator_ids)
        .exclude(round_type=3)
        .values('submission_id', 'indicator_id', 'score', 'round_type', 'id')
        .order_by('submission_id', 'indicator_id', '-round_type', '-id')
    ):
        key = (row['submission_id'], row['indicator_id'])
        if key not in review_map:
            review_map[key] = row['score']

    evidence_rows = list(
        Evidence.objects.filter(submission_id__in=submission_ids, is_deleted=False)
        .values('submission_id', 'indicator_id', 'name', 'file')
        .order_by('submission_id', 'id')
    )
    global_evidence_map = {}
    indicator_evidence_map = {}
    for ev in evidence_rows:
        submission_id = ev['submission_id']
        name = ev.get('name') or ''
        file_url = str(ev.get('file') or '')
        if ev.get('indicator_id'):
            key = (submission_id, ev['indicator_id'])
            bucket = indicator_evidence_map.setdefault(key, {'count': 0, 'names': [], 'urls': []})
            bucket['count'] += 1
            if name:
                bucket['names'].append(name)
            if file_url:
                bucket['urls'].append(file_url)
        else:
            bucket = global_evidence_map.setdefault(submission_id, {'count': 0, 'names': [], 'urls': []})
            bucket['count'] += 1
            if name:
                bucket['names'].append(name)
            if file_url:
                bucket['urls'].append(file_url)

    rows = []
    for rank, sub in enumerate(submissions, 1):
        class_obj = sub.user.class_obj
        season = sub.project.season if sub.project_id else None
        global_ev = global_evidence_map.get(sub.id, {'count': 0, 'names': [], 'urls': []})
        row = {
            'rank': rank,
            'student_no': getattr(sub.user, 'student_no', '') or sub.user.username,
            'username': sub.user.username,
            'real_name': sub.user.get_full_name() or sub.user.username,
            'gender': sub.user.get_gender_display() if hasattr(sub.user, 'get_gender_display') else '',
            'department': sub.user.department.name if sub.user.department else '',
            'major': class_obj.major.name if class_obj and class_obj.major else '',
            'class_name': class_obj.name if class_obj else '',
            'class_grade': class_obj.grade if class_obj else '',
            'phone': getattr(sub.user, 'phone', '') or '',
            'employee_no': getattr(sub.user, 'employee_no', '') or '',
            'project_status': sub.project.status if sub.project_id else '',
            'project_start_time': sub.project.start_time.isoformat() if sub.project_id and sub.project.start_time else '',
            'project_end_time': sub.project.end_time.isoformat() if sub.project_id and sub.project.end_time else '',
            'project_review_end_time': sub.project.review_end_time.isoformat() if sub.project_id and sub.project.review_end_time else '',
            'season_name': season.name if season else '',
            'season_academic_year': season.academic_year if season else '',
            'season_semester': season.semester if season else '',
            'submission_status': sub.status,
            'submitted_at': sub.submitted_at.isoformat() if sub.submitted_at else '',
            'remark': sub.remark or '',
            'final_score': float(sub.final_score) if sub.final_score is not None else 0,
            'submission_id': sub.id,
            'project_name': project.name,
            'global_evidence_count': global_ev['count'],
            'global_evidence_names': '；'.join(global_ev['names']),
            'global_evidence_urls': '\n'.join(global_ev['urls']),
        }
        for indicator_id, indicator in indicators.items():
            key = (sub.id, indicator_id)
            ans = answer_map.get(key) or {}
            imported_score = imported_map.get(key)
            arbitration_score = arbitration_map.get(key)
            reviewer_score = arbitration_score if arbitration_score is not None else review_map.get(key)
            score = get_indicator_final_score(sub, indicator)
            row[f'ind_{indicator_id}_self_score'] = float(ans['self_score']) if ans.get('self_score') is not None else ''
            row[f'ind_{indicator_id}_imported_score'] = float(imported_score) if imported_score is not None else ''
            row[f'ind_{indicator_id}_reviewer_score'] = float(reviewer_score) if reviewer_score is not None else ''
            row[f'ind_{indicator_id}_arbitration_score'] = float(arbitration_score) if arbitration_score is not None else ''
            row[f'ind_{indicator_id}_final_adopted_score'] = float(score) if score is not None else ''
            row[f'ind_{indicator_id}_process_record'] = ans.get('process_record') or ''
            ind_ev = indicator_evidence_map.get(key, {'count': 0, 'names': [], 'urls': []})
            row[f'ind_{indicator_id}_evidence_count'] = ind_ev['count']
            row[f'ind_{indicator_id}_evidence_names'] = '；'.join(ind_ev['names'])
            row[f'ind_{indicator_id}_evidence_urls'] = '\n'.join(ind_ev['urls'])
        score_detail = sub.score_detail or {}
        for parent_ind in parent_indicators:
            detail_entry = score_detail.get(str(parent_ind.id)) or {}
            score_val = detail_entry.get('score')
            row[f'agg_{parent_ind.id}'] = float(score_val) if score_val is not None else ''
        rows.append(row)
    return rows


def _field_value(row: dict, field_key: str):
    return row.get(field_key, '')


def _render_static_template_text(template: str, row: dict | None, scope_defaults: dict | None = None):
    text = str(template or '')
    row = row or {}
    scope_defaults = scope_defaults or {}

    def _replace(match):
        key = match.group(1)
        value = _field_value(row, key)
        if value in (None, ''):
            value = scope_defaults.get(key, '')
        return '' if value is None else str(value)

    return re.sub(r'\{([A-Za-z_][A-Za-z0-9_]*)\}', _replace, text)


def _display_user_name(user):
    if not user:
        return ''
    return getattr(user, 'name', '') or getattr(user, 'username', '') or ''


def normalize_group_by(group_by: str | None) -> str:
    value = (group_by or '').strip().lower()
    if value not in SUPPORTED_GROUP_BY:
        return ''
    return '' if value in ('', 'none') else value


def normalize_group_file_mode(group_file_mode: str | None, *, multi_file: bool = False) -> str:
    value = str(group_file_mode or '').strip().lower()
    if value in SUPPORTED_GROUP_FILE_MODE:
        return value
    return 'per_student_file' if multi_file else 'single_file'


def _group_label_from_row(row: dict, group_by: str) -> str:
    if group_by == 'class':
        return str(row.get('class_name') or '').strip() or '未分班'
    if group_by == 'major':
        return str(row.get('major') or '').strip() or '未分专业'
    if group_by == 'department':
        return str(row.get('department') or '').strip() or '未分院系'
    return '全部'


def _group_label_from_submission(submission, group_by: str) -> str:
    user = getattr(submission, 'user', None)
    class_obj = getattr(user, 'class_obj', None) if user else None
    if group_by == 'class':
        return getattr(class_obj, 'name', '') or '未分班'
    if group_by == 'major':
        major = getattr(class_obj, 'major', None) if class_obj else None
        return getattr(major, 'name', '') or '未分专业'
    if group_by == 'department':
        dept = getattr(user, 'department', None) if user else None
        return getattr(dept, 'name', '') or '未分院系'
    return '全部'


def split_rows_by_group(rows: list[dict], group_by: str | None) -> list[tuple[str, list[dict]]]:
    key = normalize_group_by(group_by)
    if not key:
        return [('全部', list(rows or []))]
    bucket = defaultdict(list)
    for row in rows or []:
        bucket[_group_label_from_row(row, key)].append(row)
    return sorted(bucket.items(), key=lambda item: item[0])


def _compute_scope_submission_stats(project, user, filters: dict | None, rows: list[dict], group_by: str | None = ''):
    all_scope_submissions = _resolve_submission_queryset(
        project,
        user,
        filters=filters,
        include_unscored=True,
    )
    total_all = all_scope_submissions.count()
    valid_all = len(rows or [])
    missing_all = max(total_all - valid_all, 0)
    result = {
        'scope_total_submission_count': total_all,
        'scope_valid_submission_count': valid_all,
        'scope_missing_submission_count': missing_all,
        'scope_total_submission_count_by_group': {},
        'scope_valid_submission_count_by_group': {},
        'scope_missing_submission_count_by_group': {},
    }
    key = normalize_group_by(group_by)
    if not key:
        return result
    total_by_group = defaultdict(int)
    for sub in all_scope_submissions:
        total_by_group[_group_label_from_submission(sub, key)] += 1
    valid_by_group = defaultdict(int)
    for row in rows or []:
        valid_by_group[_group_label_from_row(row, key)] += 1
    missing_by_group = {}
    for group_name in set(total_by_group.keys()) | set(valid_by_group.keys()):
        missing_by_group[group_name] = max(total_by_group.get(group_name, 0) - valid_by_group.get(group_name, 0), 0)
    result['scope_total_submission_count_by_group'] = dict(total_by_group)
    result['scope_valid_submission_count_by_group'] = dict(valid_by_group)
    result['scope_missing_submission_count_by_group'] = dict(missing_by_group)
    return result


def _normalize_reviewer_signature_policy(raw_policy: dict | None):
    policy = dict(raw_policy or {})
    source = str(policy.get('source') or 'actual_scored').strip().lower()
    if source not in {'actual_scored', 'assigned'}:
        source = 'actual_scored'
    include_arbitration = bool(policy.get('include_arbitration', False))
    role_types = policy.get('role_types') or ['counselor', 'counselor_confirm', 'counselor_dispatch', 'assistant']
    if not isinstance(role_types, list):
        role_types = ['counselor', 'counselor_confirm', 'counselor_dispatch', 'assistant']
    role_types = [str(item).strip() for item in role_types if str(item).strip()]
    assignment_statuses = policy.get('assignment_statuses') or ['assigned', 'released', 'completed']
    if not isinstance(assignment_statuses, list):
        assignment_statuses = ['assigned', 'released', 'completed']
    assignment_statuses = [str(item).strip() for item in assignment_statuses if str(item).strip()]
    return {
        'source': source,
        'include_arbitration': include_arbitration,
        'role_types': role_types,
        'assignment_statuses': assignment_statuses,
    }


def _collect_reviewer_signature_values(project, rows: list[dict], group_by: str | None = '', signature_policy: dict | None = None):
    signature_policy = _normalize_reviewer_signature_policy(signature_policy)
    submission_ids = [row.get('submission_id') for row in rows or [] if row.get('submission_id')]
    if not submission_ids:
        empty = {
            'reviewer_signatures_all': '',
            'reviewer_signatures_teachers': '',
            'reviewer_signatures_assistants': '',
        }
        return {**empty, 'by_group': {}, 'group_by': normalize_group_by(group_by)}
    assignment_qs = (
        ReviewAssignment.objects
        .filter(
            project=project,
            submission_id__in=submission_ids,
            status__in=signature_policy['assignment_statuses'],
            role_type__in=signature_policy['role_types'],
        )
        .select_related('reviewer', 'submission__user', 'submission__user__class_obj', 'submission__user__class_obj__major', 'submission__user__department')
    )
    all_names = set()
    teacher_names = set()
    assistant_names = set()
    key = normalize_group_by(group_by)
    by_group = defaultdict(lambda: {'all': set(), 'teachers': set(), 'assistants': set()})
    role_lookup = {}
    for item in assignment_qs:
        role_lookup[(item.submission_id, item.reviewer_id)] = item.role_type

    def _put_name(group_name: str, name: str, role_type: str):
        if not name:
            return
        all_names.add(name)
        if role_type == 'assistant':
            assistant_names.add(name)
        else:
            teacher_names.add(name)
        if key:
            by_group[group_name]['all'].add(name)
            if role_type == 'assistant':
                by_group[group_name]['assistants'].add(name)
            else:
                by_group[group_name]['teachers'].add(name)

    if signature_policy['source'] == 'assigned':
        for item in assignment_qs:
            group_name = _group_label_from_submission(item.submission, key) if key else '全部'
            _put_name(group_name, _display_user_name(item.reviewer), item.role_type)
    else:
        score_qs = (
            ScoreRecord.objects
            .filter(
                submission_id__in=submission_ids,
                reviewer_id__isnull=False,
                score_channel='assignment',
            )
            .exclude(round_type=3)
            .select_related('reviewer', 'submission__user', 'submission__user__class_obj', 'submission__user__class_obj__major', 'submission__user__department')
            .order_by('id')
        )
        seen_score_reviewer = set()
        for rec in score_qs:
            marker = (rec.submission_id, rec.reviewer_id)
            if marker in seen_score_reviewer:
                continue
            seen_score_reviewer.add(marker)
            role_type = role_lookup.get(marker) or ('assistant' if rec.scorer_role_level == 1 else 'counselor')
            group_name = _group_label_from_submission(rec.submission, key) if key else '全部'
            _put_name(group_name, _display_user_name(rec.reviewer), role_type)
        if signature_policy.get('include_arbitration'):
            arbitration_qs = (
                ArbitrationRecord.objects
                .filter(
                    submission_id__in=submission_ids,
                    arbitrator_id__isnull=False,
                )
                .select_related('arbitrator', 'submission__user', 'submission__user__class_obj', 'submission__user__class_obj__major', 'submission__user__department')
                .order_by('id')
            )
            seen_arbitrator = set()
            for rec in arbitration_qs:
                marker = (rec.submission_id, rec.arbitrator_id)
                if marker in seen_arbitrator:
                    continue
                seen_arbitrator.add(marker)
                group_name = _group_label_from_submission(rec.submission, key) if key else '全部'
                _put_name(group_name, _display_user_name(rec.arbitrator), 'counselor')

    def _join(values):
        return '、'.join(sorted(values))

    normalized_by_group = {}
    for group_name, payload in by_group.items():
        normalized_by_group[group_name] = {
            'reviewer_signatures_all': _join(payload['all']),
            'reviewer_signatures_teachers': _join(payload['teachers']),
            'reviewer_signatures_assistants': _join(payload['assistants']),
        }
    return {
        'reviewer_signatures_all': _join(all_names),
        'reviewer_signatures_teachers': _join(teacher_names),
        'reviewer_signatures_assistants': _join(assistant_names),
        'by_group': normalized_by_group,
        'group_by': key,
    }


def _field_numeric_values(rows: list[dict], field_key: str):
    values = []
    for row in rows or []:
        raw = row.get(field_key)
        if raw is None or raw == '':
            continue
        try:
            values.append(float(raw))
        except (TypeError, ValueError):
            continue
    return values


def _evaluate_computed_fields(rows: list[dict], computed_fields: list[dict], scope_values: dict):
    computed = {}
    for item in computed_fields or []:
        key = str(item.get('key') or '').strip()
        fn = str(item.get('fn') or '').strip().lower()
        if not key or not fn:
            continue
        target = str(item.get('target') or 'valid_rows').strip().lower()
        field_key = str(item.get('field_key') or '').strip()
        sep = str(item.get('sep') or '、')
        precision = item.get('precision')
        value = ''

        if fn == 'count':
            if target in ('missing_rows', 'missing_submissions'):
                value = int(scope_values.get('scope_missing_submission_count', 0))
            elif target == 'reviewers':
                names = str(scope_values.get('reviewer_signatures_all') or '')
                value = len([x for x in names.split('、') if x]) if names else 0
            else:
                value = len(rows or [])
        elif fn in ('sum', 'avg', 'min', 'max'):
            values = _field_numeric_values(rows, field_key)
            if not values:
                value = 0
            elif fn == 'sum':
                value = sum(values)
            elif fn == 'avg':
                value = sum(values) / len(values)
            elif fn == 'min':
                value = min(values)
            else:
                value = max(values)
            if precision is not None:
                try:
                    value = round(float(value), int(precision))
                except (TypeError, ValueError):
                    pass
        elif fn == 'join_distinct':
            if target == 'reviewers':
                role = str(item.get('role') or 'all').strip().lower()
                if role in ('teacher', 'teachers', 'counselor'):
                    value = scope_values.get('reviewer_signatures_teachers', '')
                elif role in ('assistant', 'assistants'):
                    value = scope_values.get('reviewer_signatures_assistants', '')
                else:
                    value = scope_values.get('reviewer_signatures_all', '')
            else:
                uniq = []
                seen = set()
                for row in rows or []:
                    raw = str(row.get(field_key) or '').strip()
                    if not raw or raw in seen:
                        continue
                    seen.add(raw)
                    uniq.append(raw)
                value = sep.join(uniq)
        else:
            continue
        computed[key] = value
    return computed


def enrich_export_rows(project, user, rows: list[dict], filters: dict | None = None, mapping_config: dict | None = None, group_by: str | None = ''):
    rows = list(rows or [])
    scope_stats = _compute_scope_submission_stats(project, user, filters=filters, rows=rows, group_by=group_by)
    signature_values = _collect_reviewer_signature_values(
        project,
        rows,
        group_by=group_by,
        signature_policy=(mapping_config or {}).get('reviewer_signature_policy'),
    )
    scope_values = {
        'scope_total_submission_count': scope_stats['scope_total_submission_count'],
        'scope_valid_submission_count': scope_stats['scope_valid_submission_count'],
        'scope_missing_submission_count': scope_stats['scope_missing_submission_count'],
        'reviewer_signatures_all': signature_values['reviewer_signatures_all'],
        'reviewer_signatures_teachers': signature_values['reviewer_signatures_teachers'],
        'reviewer_signatures_assistants': signature_values['reviewer_signatures_assistants'],
    }
    computed_fields = (mapping_config or {}).get('computed_fields') or []
    computed_values = _evaluate_computed_fields(rows, computed_fields, scope_values)
    scope_values.update(computed_values)

    key = signature_values.get('group_by') or normalize_group_by(group_by)
    signature_by_group = signature_values.get('by_group') or {}
    total_by_group = scope_stats.get('scope_total_submission_count_by_group') or {}
    valid_by_group = scope_stats.get('scope_valid_submission_count_by_group') or {}
    missing_by_group = scope_stats.get('scope_missing_submission_count_by_group') or {}

    for row in rows:
        row.update(scope_values)
        if key:
            group_name = _group_label_from_row(row, key)
            row['scope_total_submission_count'] = total_by_group.get(group_name, row.get('scope_total_submission_count', 0))
            row['scope_valid_submission_count'] = valid_by_group.get(group_name, row.get('scope_valid_submission_count', 0))
            row['scope_missing_submission_count'] = missing_by_group.get(group_name, row.get('scope_missing_submission_count', 0))
            sig = signature_by_group.get(group_name, {})
            if sig:
                row['reviewer_signatures_all'] = sig.get('reviewer_signatures_all', row.get('reviewer_signatures_all', ''))
                row['reviewer_signatures_teachers'] = sig.get('reviewer_signatures_teachers', row.get('reviewer_signatures_teachers', ''))
                row['reviewer_signatures_assistants'] = sig.get('reviewer_signatures_assistants', row.get('reviewer_signatures_assistants', ''))
    return rows, scope_values


def validate_export_mapping_config(mapping_config: dict | None, output_format: str) -> dict:
    """
    规范并校验导出映射配置。
    - 避免空 token、非法 Excel 列号等运行时错误。
    - 输出可直接用于渲染函数的安全配置。
    """
    cfg = dict(mapping_config or {})
    fmt = (output_format or 'xlsx').lower()

    token_mode = str(cfg.get('token_mode') or 'prefix_token').strip()
    if token_mode not in {'prefix_token', 'jinja'}:
        token_mode = 'prefix_token'
    token_prefix = str(cfg.get('token_prefix') or '@').strip() or '@'
    cfg['token_mode'] = token_mode
    cfg['token_prefix'] = token_prefix

    normalized_placeholders = []
    raw_placeholders = cfg.get('word_placeholders') or []
    if not isinstance(raw_placeholders, list):
        raise ValueError('word_placeholders 配置格式错误，应为数组')
    seen_placeholders = set()
    for idx, item in enumerate(raw_placeholders, 1):
        if not isinstance(item, dict):
            raise ValueError(f'第 {idx} 个占位符配置格式错误')
        placeholder = str(item.get('placeholder') or '').strip()
        field_key = str(item.get('field_key') or '').strip()
        if not placeholder and not field_key:
            continue
        if not placeholder or not field_key:
            raise ValueError(f'第 {idx} 个占位符配置不完整（需同时填写占位符与字段）')
        if placeholder.startswith(token_prefix):
            placeholder = placeholder[len(token_prefix):].strip()
        if not placeholder:
            raise ValueError(f'第 {idx} 个占位符名称无效')
        dedupe_key = placeholder.lower()
        if dedupe_key in seen_placeholders:
            continue
        seen_placeholders.add(dedupe_key)
        normalized_placeholders.append({
            'placeholder': placeholder,
            'field_key': field_key,
        })
    cfg['word_placeholders'] = normalized_placeholders

    normalized_columns = []
    raw_columns = cfg.get('excel_columns') or []
    if not isinstance(raw_columns, list):
        raise ValueError('excel_columns 配置格式错误，应为数组')
    seen_columns = set()
    for idx, item in enumerate(raw_columns, 1):
        if not isinstance(item, dict):
            raise ValueError(f'第 {idx} 个 Excel 列映射格式错误')
        column = str(item.get('column') or '').strip().upper()
        field_key = str(item.get('field_key') or '').strip()
        header = str(item.get('header') or '').strip()
        if not column and not field_key and not header:
            continue
        if not column or not field_key:
            raise ValueError(f'第 {idx} 个 Excel 列映射不完整（需同时填写列号与字段）')
        if not _EXCEL_COLUMN_RE.match(column):
            raise ValueError(f'第 {idx} 个 Excel 列号不合法：{column}')
        if column in seen_columns:
            raise ValueError(f'Excel 列号重复：{column}')
        seen_columns.add(column)
        normalized_columns.append({
            'column': column,
            'field_key': field_key,
            'header': header or field_key,
        })
    if fmt == 'xlsx' and not normalized_columns:
        raise ValueError('Excel 导出至少需要配置一列映射')
    cfg['excel_columns'] = normalized_columns

    try:
        header_row = int(cfg.get('header_row', 1))
        data_start_row = int(cfg.get('data_start_row', header_row + 1))
    except (TypeError, ValueError) as exc:
        raise ValueError('Excel 行号配置必须为整数') from exc
    if header_row < 1 or data_start_row < 1:
        raise ValueError('Excel 行号配置必须大于 0')
    cfg['header_row'] = header_row
    cfg['data_start_row'] = data_start_row

    # write_header: True 时在 header_row 行写入表头；False 时跳过表头写入，保留模板原有标题行
    raw_write_header = cfg.get('write_header')
    if raw_write_header is None:
        cfg['write_header'] = True
    else:
        cfg['write_header'] = bool(raw_write_header)

    # computed_fields: 声明式聚合字段（供 static_cells 与模板字段引用）
    normalized_computed = []
    raw_computed = cfg.get('computed_fields') or []
    if not isinstance(raw_computed, list):
        raise ValueError('computed_fields 配置格式错误，应为数组')
    seen_computed_keys = set()
    allowed_fns = {'count', 'sum', 'avg', 'min', 'max', 'join_distinct'}
    for idx, item in enumerate(raw_computed, 1):
        if not isinstance(item, dict):
            raise ValueError(f'第 {idx} 个 computed_fields 配置格式错误')
        key = str(item.get('key') or '').strip()
        fn = str(item.get('fn') or '').strip().lower()
        if not key and not fn:
            continue
        if not key:
            raise ValueError(f'第 {idx} 个 computed_fields 未填写 key')
        if not re.match(r'^[A-Za-z_][A-Za-z0-9_]*$', key):
            raise ValueError(f'computed_fields key 不合法：{key}')
        if key in seen_computed_keys:
            raise ValueError(f'computed_fields key 重复：{key}')
        if fn not in allowed_fns:
            raise ValueError(f'computed_fields fn 不支持：{fn}')
        target = str(item.get('target') or 'valid_rows').strip().lower() or 'valid_rows'
        field_key = str(item.get('field_key') or '').strip()
        if fn in {'sum', 'avg', 'min', 'max'} and not field_key:
            raise ValueError(f'computed_fields[{key}] 需要 field_key')
        normalized_item = {
            'key': key,
            'fn': fn,
            'target': target,
            'field_key': field_key,
            'sep': str(item.get('sep') or '、'),
        }
        if item.get('precision') is not None:
            try:
                normalized_item['precision'] = int(item.get('precision'))
            except (TypeError, ValueError) as exc:
                raise ValueError(f'computed_fields[{key}] precision 必须为整数') from exc
        role = str(item.get('role') or '').strip().lower()
        if role:
            normalized_item['role'] = role
        seen_computed_keys.add(key)
        normalized_computed.append(normalized_item)
    cfg['computed_fields'] = normalized_computed

    # static_cells: 向固定单元格写入一次的共享元数据（如专业、年级、参评人数）
    normalized_static = []
    raw_static = cfg.get('static_cells') or []
    if not isinstance(raw_static, list):
        raise ValueError('static_cells 配置格式错误，应为数组')
    seen_cells: set = set()
    for idx, item in enumerate(raw_static, 1):
        if not isinstance(item, dict):
            raise ValueError(f'第 {idx} 个静态单元格配置格式错误')
        cell = str(item.get('cell') or '').strip().upper()
        field_key = str(item.get('field_key') or '').strip()
        aggregation = str(item.get('aggregation') or 'first').strip()
        if not cell and not field_key:
            continue
        if not cell:
            raise ValueError(f'第 {idx} 个静态单元格未填写单元格地址')
        if not _CELL_RE.match(cell):
            raise ValueError(f'静态单元格地址不合法：{cell}')
        if cell in seen_cells:
            raise ValueError(f'静态单元格地址重复：{cell}')
        if aggregation not in (
            'first',
            'count',
            'first_non_empty',
            'count_valid_submissions',
            'count_missing_submissions',
            'computed',
            'template',
        ):
            aggregation = 'first'
        if aggregation == 'computed' and not field_key:
            raise ValueError(f'第 {idx} 个静态单元格使用 computed 时必须填写 field_key')
        template = str(item.get('template') or '')
        if aggregation == 'template' and not template.strip():
            raise ValueError(f'第 {idx} 个静态单元格使用 template 时必须填写模板内容')
        seen_cells.add(cell)
        normalized_static.append({
            'cell': cell,
            'field_key': field_key,
            'aggregation': aggregation,
            'template': template,
        })
    cfg['static_cells'] = normalized_static
    cfg['reviewer_signature_policy'] = _normalize_reviewer_signature_policy(cfg.get('reviewer_signature_policy'))

    return cfg


def render_excel(rows, mapping_config: dict | None = None, template_path: str | None = None) -> bytes:
    mapping_config = validate_export_mapping_config(mapping_config, output_format='xlsx')
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError('缺少 openpyxl 依赖') from exc

    if template_path:
        wb = openpyxl.load_workbook(template_path)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = '综测成绩'
    columns = mapping_config.get('excel_columns') or [
        {'column': 'A', 'field_key': 'rank', 'header': '排名'},
        {'column': 'B', 'field_key': 'student_no', 'header': '学号'},
        {'column': 'C', 'field_key': 'real_name', 'header': '姓名'},
        {'column': 'D', 'field_key': 'class_name', 'header': '班级'},
        {'column': 'E', 'field_key': 'final_score', 'header': '总分'},
    ]
    header_row = int(mapping_config.get('header_row', 1))
    data_start_row = int(mapping_config.get('data_start_row', header_row + 1))
    write_header = mapping_config.get('write_header', True)
    # Prevent header overwrite when data start row is configured too early.
    if write_header and data_start_row <= header_row:
        data_start_row = header_row + 1
    if write_header:
        for col in columns:
            ws[f"{col['column']}{header_row}"] = col.get('header') or col.get('field_key')
    # 写入静态单元格（共享元数据，只写一次）
    static_cells = mapping_config.get('static_cells') or []
    if static_cells:
        scope_defaults = {
            'scope_valid_submission_count': len(rows),
            'scope_missing_submission_count': 0,
            'scope_total_submission_count': len(rows),
        }
        for sc in static_cells:
            agg = sc.get('aggregation', 'first')
            fk = sc.get('field_key', '')
            if agg in ('count', 'count_valid_submissions') or fk in ('_count', 'scope_valid_submission_count'):
                value = len(rows)
            elif agg == 'count_missing_submissions' or fk == 'scope_missing_submission_count':
                value = _field_value(rows[0], 'scope_missing_submission_count') if rows else 0
            elif agg == 'first_non_empty':
                value = ''
                for row in rows:
                    candidate = _field_value(row, fk)
                    if str(candidate or '').strip():
                        value = candidate
                        break
            elif agg == 'computed':
                value = _field_value(rows[0], fk) if rows else ''
            elif agg == 'template':
                value = _render_static_template_text(sc.get('template', ''), rows[0] if rows else {}, scope_defaults)
            else:
                value = _field_value(rows[0], fk) if rows else scope_defaults.get(fk, '')
            ws[sc['cell']] = value
    for i, row in enumerate(rows):
        row_no = data_start_row + i
        for col in columns:
            ws[f"{col['column']}{row_no}"] = _field_value(row, col.get('field_key', ''))
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out.getvalue()


def _build_word_context(row: dict, mapping_config: dict | None = None):
    mapping_config = mapping_config or {}
    placeholders = mapping_config.get('word_placeholders') or []
    context = dict(row)
    token_prefix = mapping_config.get('token_prefix') or '@'
    for item in placeholders:
        name = (item.get('placeholder') or '').strip()
        field_key = item.get('field_key')
        if name and field_key:
            if name.startswith(token_prefix):
                name = name[len(token_prefix):]
            name = name.replace('-', '_')
            context[name] = _field_value(row, field_key)
    return context


def _replace_prefixed_tokens_in_text(text: str, token_keys, token_prefix='@'):
    if not text:
        return text
    updated = text
    for token_key in token_keys:
        token = token_key
        if token.startswith(token_prefix):
            token = token[len(token_prefix):]
        updated = updated.replace(f'{token_prefix}{token}', f'{{{{ {token.replace("-", "_")} }}}}')
    return updated


def _materialize_prefix_token_template(template_path: str, mapping_config: dict):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('缺少 python-docx 依赖') from exc
    token_prefix = mapping_config.get('token_prefix') or '@'
    token_keys = [
        (item.get('placeholder') or '').strip()
        for item in (mapping_config.get('word_placeholders') or [])
        if (item.get('placeholder') or '').strip()
    ]
    token_keys = [t for t in token_keys if t]
    if not token_keys:
        return template_path, None
    doc = Document(template_path)

    def _get_all_runs_including_hyperlinks(paragraph):
        """返回段落内所有 run（含超链接内部的 run），以及每个 run 的父元素。"""
        from docx.oxml.ns import qn
        runs = []
        # 遍历段落 XML，包括 <w:hyperlink> 内的 <w:r>
        for child in paragraph._p:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                runs.append(child)
            elif tag == 'hyperlink':
                for r in child:
                    r_tag = r.tag.split('}')[-1] if '}' in r.tag else r.tag
                    if r_tag == 'r':
                        runs.append(r)
        return runs

    def _get_run_text(r_elem):
        """从 run XML 元素中提取文本。"""
        from docx.oxml.ns import qn
        t = r_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        return t.text or '' if t is not None else ''

    def _set_run_text(r_elem, text):
        """设置 run XML 元素的文本，并保留 xml:space='preserve'。"""
        t = r_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        if t is not None:
            t.text = text
            if text and (' ' in text or text.startswith(' ') or text.endswith(' ')):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    def _replace_paragraph(paragraph):
        # 获取所有 run（含超链接内的 run）
        all_runs = _get_all_runs_including_hyperlinks(paragraph)
        if not all_runs:
            return
        # 拼接所有 run 文本整体替换，避免 Word 将 token 拆入多个 run
        full_text = ''.join(_get_run_text(r) for r in all_runs)
        replaced = _replace_prefixed_tokens_in_text(full_text, token_keys, token_prefix=token_prefix)
        if replaced != full_text:
            # 有 token 被替换：把结果写入第一个 run，清空其余 run 的文本
            _set_run_text(all_runs[0], replaced)
            for r in all_runs[1:]:
                _set_run_text(r, '')
        else:
            # 无 token：逐个 run 单独替换
            for r in all_runs:
                old = _get_run_text(r)
                new = _replace_prefixed_tokens_in_text(old, token_keys, token_prefix=token_prefix)
                if new != old:
                    _set_run_text(r, new)

    # 用 lxml 元素对象本身去重（lxml 元素支持哈希，且同一个 XML 节点
    # 的哈希值稳定，不像 id() 会在 GC 后被复用，导致不同段落误判重复）
    seen_elems: set = set()

    def _process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            elem = paragraph._p
            if elem not in seen_elems:
                seen_elems.add(elem)
                _replace_paragraph(paragraph)

    def _process_table_recursive(table):
        """递归处理表格（含嵌套表格），确保所有单元格段落都被处理。"""
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    _process_table_recursive(nested_table)

    _process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        _process_table_recursive(table)
    for section in doc.sections:
        _process_paragraphs(section.header.paragraphs)
        _process_paragraphs(section.footer.paragraphs)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_path = tmp.name
    tmp.close()
    try:
        doc.save(tmp_path)
    except Exception:
        raise
    return tmp_path, tmp_path


def _render_single_word(row: dict, mapping_config: dict, template_to_use: str) -> bytes:
    """将单个学生的数据填充到 Word 模板，返回 docx 字节。"""
    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError('缺少 docxtpl 依赖') from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError('缺少 jinja2 依赖') from exc
    tpl = DocxTemplate(template_to_use)
    context = _build_word_context(row, mapping_config)
    soft_env = Environment(undefined=_build_soft_undefined())
    tpl.render(context, jinja_env=soft_env)
    out = io.BytesIO()
    tpl.save(out)
    result = out.getvalue()
    return result


def _merge_docx_files(docx_bytes_list: list) -> bytes:
    """将多个 docx 字节合并为一个文档，每人之间插入分节符（下一页）。"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError('缺少 python-docx 依赖') from exc

    if not docx_bytes_list:
        raise RuntimeError('没有可合并的文档')
    if len(docx_bytes_list) == 1:
        return docx_bytes_list[0]

    # 以第一个文档为基础
    merged = Document(io.BytesIO(docx_bytes_list[0]))

    for extra_bytes in docx_bytes_list[1:]:
        # 在末尾段落加分节符（下一页）
        last_para = merged.paragraphs[-1] if merged.paragraphs else merged.add_paragraph()
        pPr = last_para._p.get_or_add_pPr()
        sectPr = OxmlElement('w:sectPr')
        pgSz = OxmlElement('w:pgSz')
        # 保持 A4 纵向
        pgSz.set(qn('w:w'), '11906')
        pgSz.set(qn('w:h'), '16838')
        sectPr.append(pgSz)
        pgSz_type = OxmlElement('w:type')
        pgSz_type.set(qn('w:val'), 'nextPage')
        sectPr.insert(0, pgSz_type)
        pPr.append(sectPr)

        # 追加新文档的正文内容
        extra_doc = Document(io.BytesIO(extra_bytes))
        for element in extra_doc.element.body:
            # 跳过末尾的 sectPr（文档级分节，合并时不需要）
            if element.tag.endswith('}sectPr'):
                continue
            merged.element.body.append(element)

    out = io.BytesIO()
    merged.save(out)
    return out.getvalue()


def render_word(rows, mapping_config: dict | None = None, template_path: str | None = None) -> bytes:
    """将学生数据填充到 Word 模板并返回 docx 字节。多人时合并为单一文档，每人一页。"""
    if not template_path:
        raise RuntimeError('Word 导出需要上传 Word 模板')
    mapping_config = mapping_config or {}
    token_mode = mapping_config.get('token_mode') or 'prefix_token'
    template_to_use = template_path
    cleanup_path = None
    if token_mode == 'prefix_token':
        template_to_use, cleanup_path = _materialize_prefix_token_template(template_path, mapping_config)
    try:
        if not rows:
            raise RuntimeError('没有可导出的学生数据')
        if len(rows) == 1:
            return _render_single_word(rows[0], mapping_config, template_to_use)
        # 多人：每人单独渲染后合并为一个文档
        parts = [_render_single_word(row, mapping_config, template_to_use) for row in rows]
        return _merge_docx_files(parts)
    finally:
        if cleanup_path and os.path.exists(cleanup_path):
            os.unlink(cleanup_path)


def _convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes | None:
    # Strategy 1: docx2pdf (depends on Microsoft Word on Windows/macOS)
    try:
        from docx2pdf import convert
    except ImportError:
        convert = None
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / 'input.docx'
        output_path = Path(tmpdir) / 'output.pdf'
        input_path.write_bytes(docx_bytes)
        if convert:
            try:
                convert(str(input_path), str(output_path))
            except Exception:
                pass
            else:
                if output_path.exists():
                    return output_path.read_bytes()
        # Strategy 2: LibreOffice headless convert
        soffice = _resolve_soffice_binary()
        if not soffice:
            return None
        try:
            subprocess.run(
                [
                    soffice,
                    '--headless',
                    '--convert-to',
                    'pdf',
                    '--outdir',
                    str(Path(tmpdir)),
                    str(input_path),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
        except Exception:
            return None
        if not output_path.exists():
            # Some builds may output with original filename stem
            alt_output = Path(tmpdir) / f'{input_path.stem}.pdf'
            if alt_output.exists():
                return alt_output.read_bytes()
            return None
        return output_path.read_bytes()


def _resolve_soffice_binary() -> str | None:
    """
    Resolve LibreOffice executable path across Linux/Windows environments.
    Priority:
    1) LIBREOFFICE_BIN env
    2) PATH lookup
    3) Common Windows installation paths
    """
    env_bin = os.environ.get('LIBREOFFICE_BIN', '').strip()
    if env_bin:
        expanded = os.path.expandvars(env_bin)
        if os.path.exists(expanded):
            return expanded

    for candidate in ('soffice', 'soffice.exe', 'libreoffice'):
        hit = shutil.which(candidate)
        if hit:
            return hit

    if os.name == 'nt':
        common = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        ]
        for path in common:
            if os.path.exists(path):
                return path
    return None


def _render_pdf_fallback(rows, mapping_config: dict | None = None) -> bytes:
    mapping_config = mapping_config or {}
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError('缺少 reportlab 依赖') from exc
    columns = mapping_config.get('excel_columns') or [
        {'field_key': 'rank', 'header': '排名'},
        {'field_key': 'student_no', 'header': '学号'},
        {'field_key': 'real_name', 'header': '姓名'},
        {'field_key': 'class_name', 'header': '班级'},
        {'field_key': 'final_score', 'header': '总分'},
    ]
    data = [[c.get('header') or c.get('field_key', '') for c in columns]]
    for row in rows[:1000]:
        data.append([str(_field_value(row, c.get('field_key', ''))) for c in columns])
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=landscape(A4))
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
    ]))
    doc.build([table])
    out.seek(0)
    return out.getvalue()


def render_pdf(rows, mapping_config: dict | None = None, template_path: str | None = None) -> bytes:
    # Prefer template-based rendering. If conversion runtime is unavailable,
    # keep behavior explicit so users can switch to Word export and preserve layout.
    if template_path:
        word_bytes = render_word(rows, mapping_config=mapping_config, template_path=template_path)
        pdf_bytes = _convert_docx_to_pdf_bytes(word_bytes)
        if pdf_bytes:
            return pdf_bytes
        raise RuntimeError(
            'PDF 转换失败：服务器未安装 Microsoft Word 或 LibreOffice，'
            '请改为导出 Word 格式后在本地另存为 PDF。'
        )
    # No template: directly use tabular fallback
    return _render_pdf_fallback(rows, mapping_config=mapping_config)


_FILENAME_SAFE_RE = re.compile(r'[\\/:*?"<>|\r\n\t]')


def _format_zip_filename(pattern: str, row: dict, ext: str, seen: dict) -> str:
    """
    Render a ZIP entry filename from a pattern string and a row dict.
    Supported placeholders: {student_no}, {real_name}, {username}, {class_name}, {department}, {rank}.
    Deduplicates by appending _student_no suffix on collision.
    """
    available = {
        'student_no': str(row.get('student_no') or ''),
        'real_name': str(row.get('real_name') or ''),
        'username': str(row.get('username') or ''),
        'class_name': str(row.get('class_name') or ''),
        'department': str(row.get('department') or ''),
        'rank': str(row.get('rank') or ''),
    }
    try:
        base = pattern.format(**available).strip()
    except (KeyError, ValueError):
        base = available['real_name'] or available['student_no'] or 'unknown'
    base = _FILENAME_SAFE_RE.sub('_', base).strip('_') or 'unknown'
    file_name = f'{base}.{ext}'
    if file_name in seen:
        # Disambiguate with student_no
        suffix = available['student_no'] or available['rank'] or 'dup'
        file_name = f'{base}_{suffix}.{ext}'
        # If still colliding, append a counter
        counter = 2
        original_dedup = file_name
        while file_name in seen:
            file_name = f'{original_dedup[: -(len(ext) + 1)]}_{counter}.{ext}'
            counter += 1
    seen[file_name] = True
    return file_name


def _group_dir_name_from_row(row: dict, group_by: str | None) -> str | None:
    key = normalize_group_by(group_by)
    if not key:
        return None
    if key == 'class':
        raw = row.get('class_name') or '未知班级'
    elif key == 'major':
        raw = row.get('major') or '未知专业'
    else:
        raw = row.get('department') or '未知院系'
    return _FILENAME_SAFE_RE.sub('_', str(raw)).strip() or '未分组'


def render_word_zip(rows, mapping_config: dict | None = None, template_path: str | None = None,
                    group_by: str | None = None, zip_filename_pattern: str | None = None) -> bytes:
    """每人单独渲染为一个 Word 文件，打包为 ZIP 返回。
    group_by: 'class' 时在 ZIP 内按班级建子目录；None/'' 时平铺。
    zip_filename_pattern: 文件名模板，支持 {student_no}/{real_name}/{class_name}/{department}/{rank}/{username}。
    mapping_config 中的 zip_filename_pattern 优先级高于参数。
    """
    import io
    import zipfile

    if not template_path:
        raise RuntimeError('Word 导出需要上传 Word 模板')
    if not rows:
        raise RuntimeError('没有可导出的学生数据')
    mapping_config = mapping_config or {}
    token_mode = mapping_config.get('token_mode') or 'prefix_token'
    template_to_use = template_path
    cleanup_path = None
    if token_mode == 'prefix_token':
        template_to_use, cleanup_path = _materialize_prefix_token_template(template_path, mapping_config)
    # Resolve filename pattern: mapping_config overrides function param
    pattern = mapping_config.get('zip_filename_pattern') or zip_filename_pattern or '{real_name}'
    group = (group_by or '').strip().lower()
    try:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            seen_names: dict = {}
            for row in rows:
                docx_bytes = _render_single_word(row, mapping_config, template_to_use)
                file_name = _format_zip_filename(pattern, row, 'docx', seen_names)
                dir_name = _group_dir_name_from_row(row, group)
                if dir_name:
                    zf.writestr(f'{dir_name}/{file_name}', docx_bytes)
                else:
                    zf.writestr(file_name, docx_bytes)
        return buf.getvalue()
    finally:
        if cleanup_path and os.path.exists(cleanup_path):
            os.unlink(cleanup_path)


def render_pdf_zip(rows, mapping_config: dict | None = None, template_path: str | None = None,
                   group_by: str | None = None, zip_filename_pattern: str | None = None) -> bytes:
    """每人单独渲染为一个 PDF 文件，打包为 ZIP 返回。
    group_by: 'class' 时在 ZIP 内按班级建子目录；None/'' 时平铺。
    zip_filename_pattern: 文件名模板，支持 {student_no}/{real_name}/{class_name}/{department}/{rank}/{username}。
    mapping_config 中的 zip_filename_pattern 优先级高于参数。
    """
    import io
    import zipfile

    # ZIP 分包场景优先使用模板逐人渲染；无模板时降级为逐人表格 PDF。
    if not rows:
        raise RuntimeError('没有可导出的学生数据')
    mapping_config = mapping_config or {}
    token_mode = mapping_config.get('token_mode') or 'prefix_token'
    template_to_use = template_path
    cleanup_path = None
    if token_mode == 'prefix_token':
        template_to_use, cleanup_path = _materialize_prefix_token_template(template_path, mapping_config)
    pattern = mapping_config.get('zip_filename_pattern') or zip_filename_pattern or '{real_name}'
    group = (group_by or '').strip().lower()
    try:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            seen_names: dict = {}
            for row in rows:
                if template_to_use:
                    docx_bytes = _render_single_word(row, mapping_config, template_to_use)
                    pdf_bytes = _convert_docx_to_pdf_bytes(docx_bytes)
                else:
                    pdf_bytes = None
                if not pdf_bytes:
                    if template_to_use:
                        raise RuntimeError(
                            'PDF 转换失败：服务器未安装 Microsoft Word 或 LibreOffice，'
                            '请改为导出 Word 格式后在本地另存为 PDF。'
                        )
                    # No template bound: keep per-student ZIP contract with single-row fallback table.
                    pdf_bytes = _render_pdf_fallback([row], mapping_config=mapping_config)
                file_name = _format_zip_filename(pattern, row, 'pdf', seen_names)
                dir_name = _group_dir_name_from_row(row, group)
                if dir_name:
                    zf.writestr(f'{dir_name}/{file_name}', pdf_bytes)
                else:
                    zf.writestr(file_name, pdf_bytes)
        return buf.getvalue()
    finally:
        if cleanup_path and os.path.exists(cleanup_path):
            os.unlink(cleanup_path)


def render_excel_zip_by_group(rows, mapping_config: dict | None = None, template_path: str | None = None,
                              group_by: str | None = None) -> bytes:
    import zipfile
    grouped = split_rows_by_group(rows, group_by)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        seen: dict = {}
        for group_name, group_rows in grouped:
            payload = render_excel(group_rows, mapping_config=mapping_config, template_path=template_path)
            safe_group = _FILENAME_SAFE_RE.sub('_', group_name).strip('_') or '未分组'
            file_name = _format_zip_filename(f'{safe_group}', {'real_name': safe_group}, 'xlsx', seen)
            zf.writestr(file_name, payload)
    return buf.getvalue()


def render_word_zip_by_group(rows, mapping_config: dict | None = None, template_path: str | None = None,
                             group_by: str | None = None) -> bytes:
    if not template_path:
        raise RuntimeError('Word 导出需要上传 Word 模板')
    import zipfile
    grouped = split_rows_by_group(rows, group_by)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        seen: dict = {}
        for group_name, group_rows in grouped:
            payload = render_word(group_rows, mapping_config=mapping_config, template_path=template_path)
            safe_group = _FILENAME_SAFE_RE.sub('_', group_name).strip('_') or '未分组'
            file_name = _format_zip_filename(f'{safe_group}', {'real_name': safe_group}, 'docx', seen)
            zf.writestr(file_name, payload)
    return buf.getvalue()


def render_pdf_zip_by_group(rows, mapping_config: dict | None = None, template_path: str | None = None,
                            group_by: str | None = None) -> bytes:
    import zipfile
    grouped = split_rows_by_group(rows, group_by)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        seen: dict = {}
        for group_name, group_rows in grouped:
            payload = render_pdf(group_rows, mapping_config=mapping_config, template_path=template_path)
            safe_group = _FILENAME_SAFE_RE.sub('_', group_name).strip('_') or '未分组'
            file_name = _format_zip_filename(f'{safe_group}', {'real_name': safe_group}, 'pdf', seen)
            zf.writestr(file_name, payload)
    return buf.getvalue()


def build_export_payload(project, user, filters=None, mapping_config: dict | None = None, group_by: str | None = ''):
    submissions = resolve_report_queryset(project, user, filters=filters)
    rows = build_export_rows(project, submissions)
    rows, _scope_values = enrich_export_rows(
        project,
        user,
        rows,
        filters=filters,
        mapping_config=mapping_config,
        group_by=group_by,
    )
    return rows

