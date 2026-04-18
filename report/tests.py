"""报表导出服务回归测试。"""
import io
import os
import tempfile
import zipfile

import openpyxl
from django.test import TestCase
from rest_framework.test import APIClient

from eval.models import EvalIndicator, EvalProject, EvalSeason
from org.models import Class, Department, Major
from report.models import ReportExportMapping
from report.services import (
    build_export_payload,
    get_project_export_catalog,
    render_excel,
    render_excel_zip_by_group,
    render_word,
    validate_export_mapping_config,
)
from scoring.models import ArbitrationRecord, ReviewAssignment, ScoreRecord
from submission.models import StudentSubmission
from users.models import Role, User


class ReportExportCatalogRegressionTests(TestCase):
    """覆盖仅记录模块字段口径与排序稳定性回归。"""

    def setUp(self):
        self.season = EvalSeason.objects.create(
            name='2025-2026 学年',
            academic_year='2025-2026',
            semester='1',
            status='ongoing',
        )
        self.project = EvalProject.objects.create(
            season=self.season,
            name='综合测评',
            status='ongoing',
        )

    def test_record_only_field_common_policy(self):
        """仅记录模块未推评审时：保留记录值，隐藏确认分，过程记录按开关决定。"""
        leaf = EvalIndicator.objects.create(
            project=self.project,
            name='B1 必修课门次',
            score_source='self',
            is_record_only=True,
            require_process_record=False,
            record_only_requires_review=False,
            order=1,
        )
        catalog = get_project_export_catalog(self.project, view_mode='all')
        fields = [item for item in catalog['all_fields'] if item.get('indicator_id') == leaf.id]
        key_map = {item['split_type']: item for item in fields}

        self.assertTrue(key_map['self_score']['is_common'])
        self.assertNotIn('final_adopted_score', key_map)
        self.assertFalse(key_map['process_record']['is_common'])

    def test_indicator_order_is_stable_by_module_then_indicator(self):
        """字段顺序保持 module -> indicator -> field_type 的稳定排序。"""
        EvalIndicator.objects.create(
            project=self.project,
            name='A2 项目',
            score_source='self',
            require_process_record=True,
            order=20,
        )
        EvalIndicator.objects.create(
            project=self.project,
            name='B1 项目',
            score_source='self',
            is_record_only=True,
            require_process_record=False,
            record_only_requires_review=False,
            order=30,
        )
        EvalIndicator.objects.create(
            project=self.project,
            name='A1 项目',
            score_source='self',
            require_process_record=True,
            order=10,
        )

        catalog = get_project_export_catalog(self.project, view_mode='all')
        indicator_fields = [item for item in catalog['all_fields'] if item.get('category_id') == 'indicators']
        pairs = [(item.get('module_key'), item.get('indicator_key')) for item in indicator_fields]
        first_a1 = pairs.index(('A', 'A1'))
        first_a2 = pairs.index(('A', 'A2'))
        first_b1 = pairs.index(('B', 'B1'))

        self.assertLess(first_a1, first_a2)
        self.assertLess(first_a2, first_b1)


class ReportExportCommonFieldPreferenceApiTests(TestCase):
    """覆盖“我的常用字段”接口与字段目录合并行为。"""

    def setUp(self):
        self.client = APIClient()
        self.season = EvalSeason.objects.create(
            name='2025-2026 学年',
            academic_year='2025-2026',
            semester='1',
            status='ongoing',
        )
        self.project = EvalProject.objects.create(
            season=self.season,
            name='综合测评',
            status='ongoing',
        )
        self.role = Role.objects.create(name='院系主任', code='director_for_report_test', level=3)
        self.user = User.objects.create_user(
            username='report_pref_user',
            password='Passw0rd!',
            name='报表测试用户',
            current_role=self.role,
        )
        self.client.force_authenticate(self.user)
        self.indicator = EvalIndicator.objects.create(
            project=self.project,
            name='A1 自评分',
            score_source='self',
            require_process_record=False,
            order=1,
        )
        self.process_key = f'ind_{self.indicator.id}_process_record'

    def test_common_field_preference_patch_and_read(self):
        """PATCH 增删常用字段后，GET 能正确回显。"""
        url = f'/api/v1/report/project/{self.project.id}/export/common-fields/'
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.data['common_field_keys'], [])

        resp = self.client.patch(url, {'add_keys': [self.process_key]}, format='json')
        self.assertEqual(resp.status_code, 200)
        self.assertIn(self.process_key, resp.data['common_field_keys'])

        resp = self.client.patch(url, {'remove_keys': [self.process_key]}, format='json')
        self.assertEqual(resp.status_code, 200)
        self.assertNotIn(self.process_key, resp.data['common_field_keys'])

    def test_export_fields_merge_user_common(self):
        """用户标记常用后，template_common_first 也应返回该字段。"""
        pref_url = f'/api/v1/report/project/{self.project.id}/export/common-fields/'
        fields_url = f'/api/v1/report/project/{self.project.id}/export/fields/'

        resp = self.client.patch(pref_url, {'common_field_keys': [self.process_key]}, format='json')
        self.assertEqual(resp.status_code, 200)

        resp = self.client.get(fields_url, {'view_mode': 'template_common_first'})
        self.assertEqual(resp.status_code, 200)
        field_map = {item['key']: item for item in resp.data['all_fields']}
        self.assertIn(self.process_key, field_map)
        self.assertTrue(field_map[self.process_key].get('is_user_common'))
        self.assertIn(self.process_key, resp.data.get('user_common_keys', []))

        selected_keys = {item['key'] for item in resp.data['fields']}
        self.assertIn(self.process_key, selected_keys)


class ReportExportScopeMetricsTests(TestCase):
    """覆盖静态统计字段、签名聚合与分组导出。"""

    def setUp(self):
        self.client = APIClient()
        self.season = EvalSeason.objects.create(
            name='2025-2026 学年',
            academic_year='2025-2026',
            semester='1',
            status='ongoing',
        )
        self.project = EvalProject.objects.create(
            season=self.season,
            name='综合测评',
            status='ongoing',
        )
        self.director_role = Role.objects.create(name='院系主任', code='director_export_metrics', level=3)
        self.reviewer_role = Role.objects.create(name='评审老师', code='reviewer_export_metrics', level=2)
        self.assistant_role = Role.objects.create(name='学生助理', code='assistant_export_metrics', level=1)
        self.dept = Department.objects.create(name='计算机学院', code='CS')
        self.major_a = Major.objects.create(name='计算机科学与技术', code='CS01', department=self.dept)
        self.major_b = Major.objects.create(name='软件工程', code='SE01', department=self.dept)
        self.class_a = Class.objects.create(name='计科2201', department=self.dept, major=self.major_a, grade='2022')
        self.class_b = Class.objects.create(name='软工2201', department=self.dept, major=self.major_b, grade='2022')

        self.director = User.objects.create_user(
            username='director_metrics',
            password='Passw0rd!',
            name='导出主任',
            department=self.dept,
            current_role=self.director_role,
        )
        self.reviewer = User.objects.create_user(
            username='reviewer_metrics',
            password='Passw0rd!',
            name='评审老师甲',
            department=self.dept,
            current_role=self.reviewer_role,
        )
        self.assistant = User.objects.create_user(
            username='assistant_metrics',
            password='Passw0rd!',
            name='学生助理乙',
            department=self.dept,
            current_role=self.assistant_role,
        )
        self.student_a = User.objects.create_user(
            username='student_a_metrics',
            password='Passw0rd!',
            name='学生甲',
            department=self.dept,
            class_obj=self.class_a,
            student_no='20220001',
            current_role=self.assistant_role,
        )
        self.student_b = User.objects.create_user(
            username='student_b_metrics',
            password='Passw0rd!',
            name='学生乙',
            department=self.dept,
            class_obj=self.class_a,
            student_no='20220002',
            current_role=self.assistant_role,
        )
        self.student_c = User.objects.create_user(
            username='student_c_metrics',
            password='Passw0rd!',
            name='学生丙',
            department=self.dept,
            class_obj=self.class_b,
            student_no='20220003',
            current_role=self.assistant_role,
        )
        self.sub_a = StudentSubmission.objects.create(project=self.project, user=self.student_a, status='submitted', final_score=91.5)
        self.sub_b = StudentSubmission.objects.create(project=self.project, user=self.student_b, status='submitted', final_score=None)
        self.sub_c = StudentSubmission.objects.create(project=self.project, user=self.student_c, status='submitted', final_score=88.0)
        self.indicator = EvalIndicator.objects.create(
            project=self.project,
            name='A1 基础项',
            score_source='reviewer',
            order=1,
        )
        ReviewAssignment.objects.create(
            submission=self.sub_a,
            project=self.project,
            reviewer=self.reviewer,
            role_type='counselor',
            round_type=1,
            assignment_version=1,
            status='completed',
        )
        ReviewAssignment.objects.create(
            submission=self.sub_a,
            project=self.project,
            reviewer=self.assistant,
            role_type='assistant',
            round_type=1,
            assignment_version=1,
            status='completed',
        )
        ScoreRecord.objects.create(
            submission=self.sub_a,
            indicator=self.indicator,
            reviewer=self.reviewer,
            score=90,
            round_type=1,
            score_channel='assignment',
            scorer_role_level=2,
        )
        ScoreRecord.objects.create(
            submission=self.sub_a,
            indicator=self.indicator,
            reviewer=self.assistant,
            score=89,
            round_type=1,
            score_channel='assignment',
            scorer_role_level=1,
        )

    def test_scope_metrics_and_static_cells_are_available(self):
        cfg = validate_export_mapping_config(
            {
                'excel_columns': [
                    {'column': 'A', 'field_key': 'rank', 'header': '排名'},
                    {'column': 'B', 'field_key': 'student_no', 'header': '学号'},
                ],
                'computed_fields': [
                    {'key': 'review_team_count', 'fn': 'count', 'target': 'reviewers'},
                ],
                'static_cells': [
                    {'cell': 'A1', 'field_key': 'scope_valid_submission_count', 'aggregation': 'count_valid_submissions'},
                    {'cell': 'B1', 'field_key': 'scope_missing_submission_count', 'aggregation': 'count_missing_submissions'},
                    {'cell': 'C1', 'field_key': 'reviewer_signatures_all', 'aggregation': 'first'},
                    {'cell': 'D1', 'field_key': 'review_team_count', 'aggregation': 'computed'},
                ],
            },
            output_format='xlsx',
        )
        rows = build_export_payload(
            self.project,
            self.director,
            filters={'class_id': self.class_a.id},
            mapping_config=cfg,
            group_by='class',
        )
        payload = render_excel(rows, mapping_config=cfg, template_path=None)
        wb = openpyxl.load_workbook(io.BytesIO(payload))
        ws = wb.active
        self.assertEqual(ws['A1'].value, 1)
        self.assertEqual(ws['B1'].value, 1)
        self.assertIn('评审老师甲', str(ws['C1'].value))
        self.assertIn('学生助理乙', str(ws['C1'].value))
        self.assertEqual(ws['D1'].value, 2)

    def test_reviewer_signature_policy_supports_arbitration_toggle(self):
        ArbitrationRecord.objects.create(
            submission=self.sub_a,
            indicator=self.indicator,
            arbitrator=self.director,
            score=88,
            comment='仲裁',
        )
        cfg_without_arbitration = validate_export_mapping_config(
            {
                'excel_columns': [{'column': 'A', 'field_key': 'rank', 'header': '排名'}],
                'reviewer_signature_policy': {
                    'source': 'actual_scored',
                    'include_arbitration': False,
                },
            },
            output_format='xlsx',
        )
        rows_without = build_export_payload(
            self.project,
            self.director,
            filters={'class_id': self.class_a.id},
            mapping_config=cfg_without_arbitration,
            group_by='class',
        )
        self.assertNotIn('导出主任', rows_without[0]['reviewer_signatures_all'])

        cfg_with_arbitration = validate_export_mapping_config(
            {
                'excel_columns': [{'column': 'A', 'field_key': 'rank', 'header': '排名'}],
                'reviewer_signature_policy': {
                    'source': 'actual_scored',
                    'include_arbitration': True,
                },
            },
            output_format='xlsx',
        )
        rows_with = build_export_payload(
            self.project,
            self.director,
            filters={'class_id': self.class_a.id},
            mapping_config=cfg_with_arbitration,
            group_by='class',
        )
        self.assertIn('导出主任', rows_with[0]['reviewer_signatures_all'])

    def test_static_cell_template_renders_multiple_fields(self):
        cfg = validate_export_mapping_config(
            {
                'excel_columns': [{'column': 'A', 'field_key': 'rank', 'header': '排名'}],
                'static_cells': [
                    {
                        'cell': 'A1',
                        'aggregation': 'template',
                        'template': '{major} {class_grade} 参评{scope_valid_submission_count}人',
                    },
                ],
            },
            output_format='xlsx',
        )
        rows = build_export_payload(
            self.project,
            self.director,
            filters={'class_id': self.class_a.id},
            mapping_config=cfg,
            group_by='class',
        )
        payload = render_excel(rows, mapping_config=cfg, template_path=None)
        wb = openpyxl.load_workbook(io.BytesIO(payload))
        ws = wb.active
        text = str(ws['A1'].value or '')
        self.assertIn(self.major_a.name, text)
        self.assertIn('参评1人', text)

    def test_computed_field_can_be_used_by_word_placeholder(self):
        try:
            from docx import Document
        except Exception:
            self.skipTest('python-docx 未安装，跳过 Word 占位符回归测试')

        fd, template_path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        try:
            doc = Document()
            doc.add_paragraph('参评人数：@stat_valid_count')
            doc.save(template_path)

            cfg = validate_export_mapping_config(
                {
                    'computed_fields': [
                        {'key': 'stat_valid_count', 'fn': 'count', 'target': 'valid_rows'},
                    ],
                    'word_placeholders': [
                        {'placeholder': 'stat_valid_count', 'field_key': 'stat_valid_count'},
                    ],
                },
                output_format='word',
            )
            rows = build_export_payload(
                self.project,
                self.director,
                filters={'class_id': self.class_a.id},
                mapping_config=cfg,
                group_by='class',
            )
            payload = render_word(rows, mapping_config=cfg, template_path=template_path)
            output_doc = Document(io.BytesIO(payload))
            all_text = '\n'.join(p.text for p in output_doc.paragraphs)
            self.assertIn('参评人数：1', all_text)
        finally:
            if os.path.exists(template_path):
                os.unlink(template_path)

    def test_group_context_metrics_are_isolated(self):
        rows = build_export_payload(
            self.project,
            self.director,
            filters={},
            mapping_config={'excel_columns': [{'column': 'A', 'field_key': 'rank', 'header': '排名'}]},
            group_by='major',
        )
        major_a_rows = [row for row in rows if row.get('major') == self.major_a.name]
        major_b_rows = [row for row in rows if row.get('major') == self.major_b.name]
        self.assertTrue(major_a_rows and major_b_rows)
        self.assertEqual(major_a_rows[0]['scope_valid_submission_count'], 1)
        self.assertEqual(major_a_rows[0]['scope_missing_submission_count'], 1)
        self.assertEqual(major_b_rows[0]['scope_valid_submission_count'], 1)
        self.assertEqual(major_b_rows[0]['scope_missing_submission_count'], 0)

    def test_excel_per_group_file_export_api(self):
        self.client.force_authenticate(self.director)
        mapping = ReportExportMapping.objects.create(
            name='xlsx_group_mapping',
            project=self.project,
            owner=self.director,
            output_format='xlsx',
            config={
                'excel_columns': [
                    {'column': 'A', 'field_key': 'rank', 'header': '排名'},
                    {'column': 'B', 'field_key': 'student_no', 'header': '学号'},
                ],
            },
        )
        resp = self.client.get(
            f'/api/v1/report/project/{self.project.id}/export/',
            {
                'output_format': 'xlsx',
                'mapping_id': mapping.id,
                'group_by': 'major',
                'group_file_mode': 'per_group_file',
            },
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp['Content-Type'], 'application/zip')
        zf = zipfile.ZipFile(io.BytesIO(resp.content))
        self.assertGreaterEqual(len(zf.namelist()), 2)

    def test_director_cannot_group_by_department(self):
        self.client.force_authenticate(self.director)
        mapping = ReportExportMapping.objects.create(
            name='xlsx_group_director',
            project=self.project,
            owner=self.director,
            output_format='xlsx',
            config={
                'excel_columns': [
                    {'column': 'A', 'field_key': 'rank', 'header': '排名'},
                ],
            },
        )
        resp = self.client.get(
            f'/api/v1/report/project/{self.project.id}/export/',
            {
                'output_format': 'xlsx',
                'mapping_id': mapping.id,
                'group_by': 'department',
                'group_file_mode': 'per_group_file',
            },
        )
        self.assertEqual(resp.status_code, 403)
        self.assertIn('不支持按 department 分组', str(resp.data.get('detail')))
        allowed = self.client.get(
            f'/api/v1/report/project/{self.project.id}/export/',
            {
                'output_format': 'xlsx',
                'mapping_id': mapping.id,
                'group_by': 'class',
                'group_file_mode': 'per_group_file',
            },
        )
        self.assertEqual(allowed.status_code, 200)

    def test_counselor_only_allows_class_group_by(self):
        self.client.force_authenticate(self.reviewer)
        mapping = ReportExportMapping.objects.create(
            name='xlsx_group_counselor',
            project=self.project,
            owner=self.reviewer,
            output_format='xlsx',
            config={
                'excel_columns': [
                    {'column': 'A', 'field_key': 'rank', 'header': '排名'},
                ],
            },
        )
        denied = self.client.get(
            f'/api/v1/report/project/{self.project.id}/export/',
            {
                'output_format': 'xlsx',
                'mapping_id': mapping.id,
                'group_by': 'major',
                'group_file_mode': 'per_group_file',
            },
        )
        self.assertEqual(denied.status_code, 403)
        self.assertIn('不支持按 major 分组', str(denied.data.get('detail')))

        allowed = self.client.get(
            f'/api/v1/report/project/{self.project.id}/export/',
            {
                'output_format': 'xlsx',
                'mapping_id': mapping.id,
                'group_by': 'class',
                'group_file_mode': 'per_group_file',
            },
        )
        self.assertEqual(allowed.status_code, 200)
